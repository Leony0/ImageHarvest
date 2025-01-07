from flask import Flask, render_template, request, send_from_directory,g
import sqlite3
import fitz
import os
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Mm
import cv2
import numpy as np
from PIL import Image, ImageChops

# pip insall sqlite3
# pip install opencv-python
# pip install Pillow
# pip install Flask
# pip install python-docx
# pip install PyMuPDF

app = Flask(__name__)



def get_db():
    if 'db' not in g:
        # データベースをオープンしてFlaskのグローバル変数に保存
        g.db = sqlite3.connect('login_DB.db')
    return g.db

@app.route('/',methods=['GET','POST'])
def new():
    if request.method=="POST":
    # ログインに使う名前とパスワードの取得
        name = request.form["name"]
        password = request.form["password"]
      
        # データベースを開く
        con = get_db()

        # データベースとログイン情報を参照する
        cur = con.execute("select * from 登録一覧")
        for i in cur:
            # 名前とパスワードが一致すればログインできる
            if name ==i[1] and password==i[2]:
                cur = con.execute("select * from 登録一覧 order by コード")
                data = cur.fetchall()
                con.close()
                
                return render_template('index.html', data = data)
        cur.close()
        return render_template('login.html',error_message="ログインできませんでした")


    else:
        # データベースを開く
        con = get_db()

        # テーブルの有無を確認する
        cur = con.execute("select count(*) from sqlite_master where TYPE='table' AND name='登録一覧'")

        for row in cur:
            if row[0] == 0:
                # テーブルを作成する（コードにより管理）
                cur.execute("CREATE TABLE 登録一覧(コード INTEGER PRIMARY KEY, 名前 STRING, パスワード STRING)")
                # レコードを作成する
                # 例として，同志社太郎のユーザー情報を生成
                cur.execute(
                    """INSERT INTO 登録一覧(コード, 名前, パスワード) 
                    values(1, '同志社太郎', 'p443')
                
                    """)
                con.commit()
        
        # 登録一覧を読み込む
        cur = con.execute("select * from 登録一覧 order by コード")
        data = cur.fetchall()
        con.close()

        return render_template('login.html')
    
#　home画面の表示
@app.route('/index', methods=["GET","POST"])
def index():
    if request.method=="POST":
        return render_template('index.html')
    return render_template('index.html')

@app.route('/new', methods=["GET","POST"])
def new_post():
    if request.method=="POST":
        # 新規登録する名前とパスワードの取得
        name = request.form["name"]
        password = request.form["password"]
        password=str(password)

        # データベースを開く
        con = get_db()

        # コードは既に登録されているコードの最大値＋１の値で新規登録を行う
        cur = con.execute("select MAX(コード) AS max_code from 登録一覧")
        for row in cur:
            new_code = row[0] + 1
        cur.close()

        # 登録処理
        sql = "INSERT INTO 登録一覧(コード, 名前, パスワード)values({},'{}','{}')".format(new_code, name, password)
        con.execute(sql)
        con.commit()

        # 登録一覧を読み込む
        cur = con.execute("select * from 登録一覧 order by コード")
        data = cur.fetchall()
        con.close()
        
        message="登録できました"
        return render_template('new.html',message=message)
        
    else:
        return render_template('new.html')

# アプリの概要ページに遷移
@app.route('/about')
def about():
    return render_template('about.html')

#############################################################################################################################################
#############################################################################################################################################

#　PDF→画像抽出アプリについて
@app.route('/app1', methods=["GET","POST"])
def app1():
    # 選択肢の表示
    dir_path = Path("./static/image/app1_image")
    filelist_img = os.listdir(dir_path)
    pdf_path = Path("./files/app1_files")
    filelist_pdf = os.listdir(pdf_path)

    if request.method=="POST":
        return render_template('app1.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    return render_template('app1.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf)

# ファイルのアップロード＋画像抽出＋wordダウンロードまで一度に
@app.route('/upload_extract', methods=["POST"])
def upload_extract():
    if 'file' not in request.files:
            error_message1="ファイルが選択されていません"

            dir_path = Path("./static/image/app1_image")
            filelist_img = os.listdir(dir_path)
            pdf_path = Path("./files/app1_files")
            filelist_pdf = os.listdir(pdf_path)

            return render_template('app1.html',error_message1=error_message1, filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    file = request.files.get('file')
    if file.filename == '':
            error_message1="ファイルが選択されていません"
            
            dir_path = Path("./static/image/app1_image")
            filelist_img = os.listdir(dir_path)
            pdf_path = Path("./files/app1_files")
            filelist_pdf = os.listdir(pdf_path)

            return render_template('app1.html',error_message1=error_message1, filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    
    # 入力PDFのアップロード
    file_name = file.filename
    file_path = os.path.join(os.getcwd(),'files',"app1_files", file_name)
    file.save(file_path)
    
    # 画像抽出について

    # 画像の保存先フォルダを設定
    filename = file_name
    dir_name = filename.split('.')[0]
    img_dir = os.path.join(os.getcwd(),"static","image","app1_image",dir_name) #階層の一個下げ
    if os.path.isdir(img_dir) == False:
        os.mkdir(img_dir)    
    
    #　PDFファイルの読み込み
    doc = fitz.open(Path("./files/app1_files/"+filename))

    # 画像情報を格納するリスト
    images = []   

    # １ページずつ画像データを取得
    for page in range(len(doc)):
        images.append(doc[page].get_images())
 
    # ページ内の画像情報を順番に処理
    for pageNo, image in enumerate(images):
    # ページ内の画像情報を処理する
        if image != []:
            for i in range(len(image)):
                # 画像情報の取得
                xref = image[i][0]
                smask = image[i][1]
                if image[i][8] == 'FlateDecode':
                    ext = 'png'
                elif image[i][8] == 'DCTDecode':
                    ext = 'jpeg'
 
                # マスク情報の取得と画像の再構築
                pix = fitz.Pixmap(doc.extract_image(xref)["image"])
                if smask > 0:
                    mask = fitz.Pixmap(doc.extract_image(smask)["image"])
                    pix = fitz.Pixmap(pix, 0) 
                    pix = fitz.Pixmap(pix, mask)
                # 画像を保存
                img_name = os.path.join(img_dir, f'image{pageNo+1}_{i}.{ext}')
                pix.save(img_name)

    #wordへの書き出しについて

    # Wordファイルを開く
    doc = Document()

    # 画像フォルダのパスを設定する
    folder_path = Path("./static/image/app1_image/"+dir_name)

    # フォルダ内のすべてのPNGファイルとJPGファイルを取得
    image_files = folder_path.glob("*")  

    for file in image_files:
    # 画像の挿入
        doc.add_picture(str(file), width=Mm(80))

    # Wordファイルを保存する
    folder_path = Path("./word_files/app1_word")
    word_path = os.path.join(folder_path, f'{dir_name}.docx')
    doc.save(word_path)   

    # パスの取得

    dir_path = Path("./static/image/app1_image")
    filelist_img = os.listdir(dir_path)
    pdf_path = Path("./files/app1_files")
    filelist_pdf = os.listdir(pdf_path)

    # img_staticpath = "./static/image/app1_image/sample"
    img_dir = Path("./static/image/app1_image/"+dir_name)
    
    imglist = os.listdir(img_dir)

    word_message=dir_name+".docxに貼り付けされました"
    docx_name=dir_name+".docx"
    docx_path=folder_path / docx_name
    print(docx_path)

    return render_template('app1.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf, img_dir=img_dir, imglist=imglist,word_message=word_message, docx_path=docx_path, docx_name=docx_name)

# wordのダウンロード機能について
@app.route('/word_files/<path:filename>')
def download_file(filename):
    return send_from_directory(directory='word_files', path=filename)

# ファイルの削除を行う
@app.route('/delete_pdf', methods=["POST"])
def delete_pdf():
    target_pdf = request.form.get('mySelect_pdf')
    if target_pdf == '':
            error_message2="ファイルが選択されていません"

            dir_path = "./static/image/app1_image"
            filelist_img = os.listdir(dir_path)
            pdf_path = "./files/app1_files"
            filelist_pdf = os.listdir(pdf_path)

            return render_template('app1.html',error_message2=error_message2, filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    
    os.remove("./files/app1_files/"+target_pdf)

    dir_path = "./static/image/app1_image"
    filelist_img = os.listdir(dir_path)
    pdf_path = "./files/app1_files"
    filelist_pdf = os.listdir(pdf_path)

    return render_template('app1.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf)

# 画像データの削除
@app.route('/delete_data', methods=["POST"])
def delete_data():
    target_dir = request.form.get('mySelect_img')
    if target_dir == '':
            error_message3="ファイルが選択されていません"

            dir_path = "./static/image/app1_image"
            filelist_img = os.listdir(dir_path)
            pdf_path = "./files/app1_files"
            filelist_pdf = os.listdir(pdf_path)

            return render_template('app1.html',error_message3=error_message3, filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    shutil.rmtree("./static/image/app1_image/"+target_dir)

    dir_path = "./static/image/app1_image"
    filelist_img = os.listdir(dir_path)
    pdf_path = "./files/app1_files"
    filelist_pdf = os.listdir(pdf_path)

    return render_template('app1.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf)


#############################################################################################################################################
#############################################################################################################################################

#　画像トリミングについて
@app.route('/app2', methods=['GET', 'POST'])
def app2():
    
    dir_path = Path("./static/image/app2_image")
    filelist_img = os.listdir(dir_path)
    pdf_path = Path("./files/app2_files")
    filelist_pdf = os.listdir(pdf_path)

    
    if request.method == 'POST':
        
        # ファイルがアップロードされた場合
        if 'file' not in request.files:
            error_message1="ファイルが選択されていません"
            return render_template('app2.html',error_message1=error_message1,filelist_img=filelist_img, filelist_pdf=filelist_pdf)
        file = request.files['file']
        if file.filename == '':
            error_message1="ファイルが選択されていません"
            return render_template('app2.html',error_message1=error_message1,filelist_img=filelist_img, filelist_pdf=filelist_pdf)
        if file:
            # ファイルを保存
            folder_path = Path("./files/app2_files")
            file_path = os.path.join(folder_path, file.filename)
            file.save(file_path)
            return render_template('app2.html', filename=file.filename,filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    return render_template('app2.html', filename=None,filelist_img=filelist_img, filelist_pdf=filelist_pdf)

# ファイルのアップロード
@app.route('/uploads/<filename>')
def uploaded_file(filename):
    folder_path = Path("./files/app2_files")
    return send_from_directory(folder_path, filename)

#　画像トリミング機能について（javascriptと連動）
@app.route('/cut_image', methods=['POST'])
def cut_image():
    data = request.get_json()
    filename = data['filename']
    x1, y1 = data['x1'], data['y1']
    #print( data['x1'])
    #print( data['y1'])
    x2, y2 = data['x2'], data['y2']
    #print( data['x2'])
    #print( data['y2'])

    folder_path = Path("./files/app2_files")
    file_path = os.path.join(folder_path, filename)
    img = cv2.imread(file_path, cv2.IMREAD_COLOR)

    if x1 > x2:
        x1, x2 = x2, x1
    if y1 > y2:
        y1, y2 = y2, y1

    #画面調節するためのコード
    #1000はテキトーな値にとりあえず設定
    h, w, _ = img.shape
    rate=1000/w
    
    x1=int(x1/rate)
    x2=int(x2/rate)
    y1=int(y1/rate)
    y2=int(y2/rate)
    print(w, h, x1, x2, y1, y2)

    
    cropped_img = img[y1:y2, x1:x2]


    #ファイルへの書き出し
    dir_name = filename.split('.')[0]
    image_folder_path = Path("./static/image/app2_image/"+dir_name)
    if os.path.isdir(image_folder_path) == False:
        os.mkdir(image_folder_path) 

    new_filename = 'new_file{:03d}.jpg'.format(len(os.listdir(image_folder_path)))
    new_file_path = os.path.join(image_folder_path, new_filename)
    cv2.imwrite(new_file_path, cropped_img)

             

    return {'new_filename': new_filename}

# 入力ファイルの削除を行う
@app.route('/delete_pdf2', methods=["POST"])
def delete_pdf2():
    
    target_pdf = request.form.get('mySelect_pdf')
    if target_pdf == '':
            error_message2="ファイルが選択されていません"

            target_pdf = request.form.get('mySelect_pdf')
            dir_path = "./static/image/app2_image"
            filelist_img = os.listdir(dir_path)
            pdf_path = "./files/app2_files"
            filelist_pdf = os.listdir(pdf_path)
            
            return render_template('app2.html',error_message2=error_message2, filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    #削除機能
    os.remove("./files/app2_files/"+target_pdf)

    dir_path = "./static/image/app2_image"
    filelist_img = os.listdir(dir_path)
    pdf_path = "./files/app2_files"
    filelist_pdf = os.listdir(pdf_path)

    return render_template('app2.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf)

#画像データの削除
@app.route('/delete_data2', methods=["POST"])
def delete_data2():

    target_dir = request.form.get('mySelect_img')
    if target_dir == '':
            error_message3="ファイルが選択されていません"
                
            dir_path = "./static/image/app2_image"
            filelist_img = os.listdir(dir_path)
            pdf_path = "./files/app2_files"
            filelist_pdf = os.listdir(pdf_path)

            return render_template('app2.html',error_message3=error_message3, filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    #フォルダの削除
    shutil.rmtree("./static/image/app2_image/"+target_dir)
    
    dir_path = "./static/image/app2_image"
    filelist_img = os.listdir(dir_path)
    pdf_path = "./files/app2_files"
    filelist_pdf = os.listdir(pdf_path)

    return render_template('app2.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf)

#wordダウンロード機能
@app.route('/download2', methods=["POST"])
def download2():
    target_dir = request.form.get('mySelect_word')
    if target_dir == '':
            error_message3="ファイルが選択されていません"
                
            dir_path = "./static/image/app2_image"
            filelist_img = os.listdir(dir_path)
            pdf_path = "./files/app2_files"
            filelist_pdf = os.listdir(pdf_path)

            return render_template('app2.html',error_message3=error_message3, filelist_img=filelist_img, filelist_pdf=filelist_pdf)

    #wordへの書き出し

    # Wordファイルを開く
    doc = Document()

    # 画像フォルダのパスを設定する
    folder_path = Path("./static/image/app2_image/"+target_dir)

    # フォルダ内のすべてのPNGファイルとJPGファイルを取得
    image_files = folder_path.glob("*")  

    for file in image_files:
    # 画像の挿入
        doc.add_picture(str(file), width=Mm(80))

    # Wordファイルを保存する
    folder_path = Path("./word_files/app2_word")
    word_path = os.path.join(folder_path, f'{target_dir}.docx')
    doc.save(word_path) 
    
    word_message=target_dir+".docxに貼り付けされました"
    docx_name=target_dir+".docx"
    docx_path=folder_path / docx_name
    print(docx_path)

    dir_path = "./static/image/app2_image"
    filelist_img = os.listdir(dir_path)
    pdf_path = "./files/app2_files"
    filelist_pdf = os.listdir(pdf_path)
    return render_template('app2.html',word_message=word_message, docx_path=docx_path, docx_name=docx_name, filelist_img=filelist_img, filelist_pdf=filelist_pdf)

##############################################################################################################################################################
##############################################################################################################################################################

# スキャンされた画像から自動抽出
@app.route('/app3', methods=["GET","POST"])
def app3():
    dir_path = "./static/image/app3_image"
    filelist_img = os.listdir(dir_path)
    pdf_path = "./files/app3_files"
    filelist_pdf = os.listdir(pdf_path)
    if request.method=="POST":
        return render_template('app3.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    return render_template('app3.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf)

@app.route('/upload3', methods=["POST"])
def upload3():
    #アップロード
    if 'file' not in request.files:
            error_message1="ファイルが選択されていません"
                
            dir_path = "./static/image/app3_image"
            filelist_img = os.listdir(dir_path)
            pdf_path = "./files/app3_files"
            filelist_pdf = os.listdir(pdf_path)

            return render_template('app3.html',error_message1=error_message1, filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    file = request.files.get('file')
    if file.filename == '':
            error_message1="ファイルが選択されていません"
                
            dir_path = "./static/image/app3_image"
            filelist_img = os.listdir(dir_path)
            pdf_path = "./files/app3_files"
            filelist_pdf = os.listdir(pdf_path)

            return render_template('app3.html',error_message1=error_message1, filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    file_name = file.filename
    file_path = os.path.join(os.getcwd(),'files',"app3_files", file_name)
    file.save(file_path)
    

    #画像抽出
    # 画像を読み込む
    path = Path("./files/app3_files")
    file_path = os.path.join(path, file_name)
    img = cv2.imread(file_path)
    # HSV 色空間に変換
    hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
    # 2値化で赤の枠線を抽出
    #binary = cv2.inRange(hsv, (170, 128, 0), (180, 255, 255))

    # 赤色の範囲を設定
    lower_red1 = np.array([0, 100, 100])
    upper_red1 = np.array([10, 255, 255])
    lower_red2 = np.array([170, 100, 100])
    upper_red2 = np.array([180, 255, 255])

    # 赤色の範囲に対してマスクを作成
    mask1 = cv2.inRange(hsv, lower_red1, upper_red1)
    mask2 = cv2.inRange(hsv, lower_red2, upper_red2)
    mask = cv2.bitwise_or(mask1, mask2)

    # モルフォロジー演算でノイズを除去
    kernel = np.ones((3, 3), np.uint8)
    eroded = cv2.morphologyEx(mask, cv2.MORPH_CLOSE, kernel)
    eroded = cv2.morphologyEx(mask, cv2.MORPH_OPEN, kernel)

    # 輪郭を抽出
    contours, hierarchy = cv2.findContours(
        eroded, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE
    )
    # 誤検出の輪郭を消去
    contours = list(filter(lambda x: cv2.contourArea(x) > 100, contours))

    # 輪郭の凸包に置き換える
    convex_hulls = list(map(cv2.convexHull, contours))

    # 検出された輪郭内部を (255, 255, 255) で塗りつぶす
    mask = np.zeros_like(img)
    cv2.drawContours(mask, convex_hulls, -1, color=(255, 255, 255), thickness=-1)

    # 元画像でマスクの値が (255, 255, 255) の画素以外を白に置換する
    img = np.where(mask == 255, img, 255)
    # 結果を保存する

    dir_name = file_name.split('.')[0]
    image_folder_path = Path("./static/image/app3_image/"+dir_name)
    if os.path.isdir(image_folder_path) == False:
        os.mkdir(image_folder_path)    
    new_filename = 'new_file000.jpg'.format(len(os.listdir(image_folder_path)))
    new_file_path = os.path.join(image_folder_path, new_filename)
    cv2.imwrite(new_file_path, img)
    ##ここまでで，余白画像を作成（赤枠内部だけ画像があり，それ以外の部分が真っ白の画像）

    #トリミング実装
    image = Image.open(new_file_path)
    # getpixel(0, 0) で左上の色を取得し、背景色のみの画像を作成する
    bg = Image.new(image.mode, image.size, image.getpixel((0, 0)))
    # 背景色画像と元画像の差分を取得
    diff = ImageChops.difference(image, bg)
    #diff.show()
    diff = ImageChops.add(diff, diff, 2.0, -100)
    #diff.show()

    # 差分画像を二値化
    diff_np = np.array(diff)
    gray = cv2.cvtColor(diff_np, cv2.COLOR_BGR2GRAY)
    _, binary = cv2.threshold(gray, 1, 255, cv2.THRESH_BINARY)

    # 連結したコンポーネントを検出
    contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    # 各コンポーネント（物体）のトリミング
    # 余白を作って切り出している
    trimmed_images = []
    offset=10
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        bbox = (x - offset, y - offset, x + w + offset, y + h + offset)
        bbox = (max(bbox[0], 0), max(bbox[1], 0), min(bbox[2], image.width), min(bbox[3], image.height))
        cropped = image.crop(bbox)
        trimmed_images.append(cropped)

    # トリミングした画像を保存
    #　まずフォルダの作成
    #dir_name = file_name.split('.')[0]
    #image_folder_path = Path("./image/app3_image/"+dir_name)
    #if os.path.isdir(image_folder_path) == False:
    #    os.mkdir(image_folder_path)    
    #new_filename = 'new_file{:03d}.jpg'.format(len(os.listdir(image_folder_path)))
    #new_file_path = os.path.join(image_folder_path, new_filename)
    #cv2.imwrite(new_file_path, img)

    for i, trimmed_image in enumerate(trimmed_images):
        img_name = os.path.join(image_folder_path, f'trimmed_object_{i}.jpg')
        trimmed_image.save(img_name)

    # image_folder_path = Path(image_folder_path + f"/new_file000.jpg")
    # print(image_folder_path)
    # os.remove("./files/app3_files/"+target_pdf)

    # 余白を作成した画像消す必要あり（プレビュー，ワードに貼り付けされてしまうため）
   # print("パス"+new_file_path)
    os.remove(new_file_path)

#wordへの書き出し


    # Wordファイルを開く
    doc = Document()

    # 画像フォルダのパスを設定する
    folder_path = Path("./static/image/app3_image/"+dir_name)

    # フォルダ内のすべてのPNGファイルとJPGファイルを取得
    image_files = folder_path.glob("*")  

    for file in image_files:
    # 画像の挿入
        doc.add_picture(str(file), width=Mm(80))

    # Wordファイルを保存する
    folder_path = Path("./word_files/app3_word")
    word_path = os.path.join(folder_path, f'{dir_name}.docx')
    doc.save(word_path)

    dir_path = "./static/image/app3_image"
    filelist_img = os.listdir(dir_path)
    pdf_path = "./files/app3_files"
    filelist_pdf = os.listdir(pdf_path)            

    imglist = os.listdir(image_folder_path)
    print(image_folder_path)

    word_message=dir_name+".docxに貼り付けされました"
    docx_name=dir_name+".docx"
    docx_path=folder_path / docx_name
    print(docx_path)

    return render_template('app3.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf, image_folder_path=image_folder_path, imglist=imglist,word_message=word_message, docx_path=docx_path, docx_name=docx_name)


# ファイルの削除を行う

@app.route('/delete_pdf3', methods=["POST"])
def delete_pdf3():
    target_pdf = request.form.get('mySelect_pdf')
    if target_pdf == '':
            error_message2="ファイルが選択されていません"
                
            dir_path = "./static/image/app3_image"
            filelist_img = os.listdir(dir_path)
            pdf_path = "./files/app3_files"
            filelist_pdf = os.listdir(pdf_path)

            return render_template('app3.html',error_message2=error_message2, filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    os.remove("./files/app3_files/"+target_pdf)
    
    dir_path = "./static/image/app3_image"
    filelist_img = os.listdir(dir_path)
    pdf_path = "./files/app3_files"
    filelist_pdf = os.listdir(pdf_path)

    return render_template('app3.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf)

#画像の削除
@app.route('/delete_data3', methods=["POST"])
def delete_data3():
    target_dir = request.form.get('mySelect_img')
    if target_dir == '':
            error_message3="ファイルが選択されていません"
                
            dir_path = "./static/image/app3_image"
            filelist_img = os.listdir(dir_path)
            pdf_path = "./files/app3_files"
            filelist_pdf = os.listdir(pdf_path)

            return render_template('app3.html',error_message3=error_message3, filelist_img=filelist_img, filelist_pdf=filelist_pdf)
    shutil.rmtree("./static/image/app3_image/"+target_dir)
    
    dir_path = "./static/image/app3_image"
    filelist_img = os.listdir(dir_path)
    pdf_path = "./files/app3_files"
    filelist_pdf = os.listdir(pdf_path)

    return render_template('app3.html', filelist_img=filelist_img, filelist_pdf=filelist_pdf)

    


if __name__ == "__main__":
    app.run(debug=True, port=8888, threaded=True)  
